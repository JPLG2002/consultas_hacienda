"""
Script FINAL para crear GUIA_USUARIO_SISTEMA profesional
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_professional_guide():
    # Rutas
    original_path = r"C:\Users\15pab\Desktop\consultas_hacienda\consultas_hacienda\GUIA_USUARIO_SISTEMA.docx"
    output_path = r"C:\Users\15pab\Desktop\consultas_hacienda\consultas_hacienda\GUIA_DE_USUARIO_OPERADORES.docx"
    
    print("Creando nuevo documento profesional...")
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========== PORTADA ==========
    print("Creando portada profesional...")
    
    # Espaciado superior
    for _ in range(2):
        doc.add_paragraph()
    
    # Título institucional
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ALCALDÍA DE SANTIAGO DE CALI")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Secretaría de Desarrollo Institucional")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Espacio
    for _ in range(3):
        doc.add_paragraph()
    
    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUÍA DE USUARIO")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Consulta de Actos Administrativos")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(51, 102, 153)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Departamento Administrativo de Hacienda")
    run.font.size = Pt(14)
    run.italic = True
    
    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Espacio
    for _ in range(4):
        doc.add_paragraph()
    
    # Informe técnico
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Informe Técnico Complementario:")
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Diseño y Funcionamiento del Sistema de Consulta de Hacienda")
    run.font.size = Pt(12)
    
    # Espacio
    for _ in range(3):
        doc.add_paragraph()
    
    # Responsables
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RESPONSABLES DEL DESARROLLO")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Juan Pablo Lagos Gómez")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Juan Felipe Ramírez")
    run.font.size = Pt(12)
    
    # Espacio
    for _ in range(3):
        doc.add_paragraph()
    
    # Fecha
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Santiago de Cali, 28 de noviembre de 2025")
    run.font.size = Pt(11)
    run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión 1.0")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Salto de página
    doc.add_page_break()
    
    # ========== TABLA DE CONTENIDO ==========
    print("Creando tabla de contenido...")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TABLA DE CONTENIDO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    toc_items = [
        "1. Introducción al Sistema",
        "2. Acceso al Sistema - Inicio de Sesión",
        "3. Módulo de Carga de Archivos",
        "4. Mis Registros",
        "5. Auditoría del Sistema",
        "6. Gestión de Usuarios",
        "7. Consulta Pública de Actos Administrativos",
        "8. Roles de Usuario y sus Diferencias",
        "9. Subida de PDFs vía FTP",
        "10. Preguntas Frecuentes"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ========== INTRODUCCIÓN ==========
    print("Creando sección de introducción...")
    
    p = doc.add_paragraph()
    run = p.add_run("1. INTRODUCCIÓN AL SISTEMA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    intro_text = """El Sistema de Consulta de Actos Administrativos es una herramienta web desarrollada para el Departamento Administrativo de Hacienda de la Alcaldía de Santiago de Cali. Esta guía proporciona instrucciones detalladas para el uso correcto del sistema por parte de los operadores y demás usuarios autorizados.

El sistema ha sido diseñado para cumplir con los siguientes objetivos:

• Facilitar la carga masiva de actos administrativos mediante archivos Excel
• Permitir la consulta pública de actos administrativos por parte de los ciudadanos
• Garantizar la trazabilidad completa de todas las operaciones mediante auditoría
• Administrar usuarios con diferentes niveles de permisos según su rol institucional
• Asociar documentos PDF a cada acto administrativo para consulta pública"""
    
    p = doc.add_paragraph()
    run = p.add_run(intro_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== ACCESO AL SISTEMA ==========
    print("Creando sección de acceso al sistema...")
    
    p = doc.add_paragraph()
    run = p.add_run("2. ACCESO AL SISTEMA - INICIO DE SESIÓN")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    login_text = """Para acceder al sistema, siga estos pasos:

1. Abra su navegador web (Chrome, Firefox o Edge recomendados)
2. Ingrese la URL del sistema proporcionada por el administrador
3. En la pantalla de inicio de sesión, ingrese su usuario y contraseña
4. Haga clic en el botón "Iniciar Sesión"

Si tiene problemas con su acceso, comuníquese con el administrador del sistema de su dependencia.

IMPORTANTE: No comparta sus credenciales de acceso con terceros. Cada usuario es responsable de las acciones realizadas con su cuenta."""
    
    p = doc.add_paragraph()
    run = p.add_run(login_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== MÓDULO DE CARGA ==========
    print("Creando sección de carga de archivos...")
    
    p = doc.add_paragraph()
    run = p.add_run("3. MÓDULO DE CARGA DE ARCHIVOS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    carga_text = """Esta pestaña permite subir archivos Excel con datos de actos administrativos para cargarlos masivamente al sistema.

FUNCIONALIDADES:
• Subir archivos Excel (.xlsx, .xls) con los datos de los actos administrativos
• El sistema valida que el formato del Excel sea correcto
• Verifica que existan los PDFs asociados a cada registro
• Muestra archivos pendientes de procesar
• Permite procesar los archivos para insertar los datos en la base de datos

PROCESO DE CARGA:
1. Prepare su archivo Excel con el formato requerido (todas las columnas obligatorias)
2. Suba los PDFs asociados al servidor vía FTP
3. Haga clic en "Subir Archivo Excel" y seleccione su archivo
4. Verifique que los PDFs estén disponibles usando "Verificar PDFs"
5. Procese el archivo para cargar los datos a la base de datos

NOTA: Solo disponible para roles admin_general, admin_dependencia y operador."""
    
    p = doc.add_paragraph()
    run = p.add_run(carga_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== MIS REGISTROS ==========
    print("Creando sección de mis registros...")
    
    p = doc.add_paragraph()
    run = p.add_run("4. MIS REGISTROS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    registros_text = """Esta pestaña muestra todos los registros que el usuario ha cargado al sistema.

FUNCIONALIDADES:
• Lista los actos administrativos cargados por el usuario logueado
• Permite ver el estado de cada registro (en trámite o finalizado)
• Opción de eliminar registros (borrado lógico) si el usuario tiene permisos
• Restaurar registros previamente eliminados
• Los administradores pueden ver registros de otros usuarios según su nivel de permisos

ESTADOS DE LOS REGISTROS:
• En Trámite: El acto administrativo está vigente y la fecha de desfijación no ha pasado
• Histórico: La fecha de desfijación ya pasó y el registro se considera archivado

PERMISOS DE ELIMINACIÓN:
Si tiene habilitados los permisos de eliminación, verá un botón para eliminar cada registro. La eliminación es "lógica", lo que significa que el registro se oculta pero puede ser restaurado por un administrador."""
    
    p = doc.add_paragraph()
    run = p.add_run(registros_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== AUDITORÍA ==========
    print("Creando sección de auditoría...")
    
    p = doc.add_paragraph()
    run = p.add_run("5. AUDITORÍA DEL SISTEMA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    auditoria_text = """Esta pestaña registra todas las acciones realizadas en el sistema para trazabilidad y seguridad.

FUNCIONALIDADES:
• Registro de inicios y cierres de sesión
• Historial de subidas de archivos Excel
• Procesamiento de datos (exitosos, parciales o fallidos)
• Eliminación y restauración de registros
• Filtros por usuario, acción y fecha (solo para administradores)

NIVELES DE ACCESO A LA AUDITORÍA:
• Admin General: Ve toda la auditoría con filtros completos
• Admin Dependencia: Ve auditoría de usuarios de su dependencia
• Operador: Solo ve su propia actividad
• Consulta: No tiene acceso a esta pestaña

ACCIONES REGISTRADAS:
• login/logout: Inicio y cierre de sesión
• upload_excel: Subida de archivo Excel
• procesar_exito: Procesamiento exitoso
• borrar_registro: Eliminación de registro
• restaurar_registro: Restauración de registro eliminado"""
    
    p = doc.add_paragraph()
    run = p.add_run(auditoria_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== GESTIÓN DE USUARIOS ==========
    print("Creando sección de gestión de usuarios...")
    
    p = doc.add_paragraph()
    run = p.add_run("6. GESTIÓN DE USUARIOS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    usuarios_text = """Esta pestaña permite administrar los usuarios del sistema (solo visible para administradores).

FUNCIONALIDADES:
• Crear nuevos usuarios con diferentes roles
• Asignar permisos de eliminación de registros
• Activar/desactivar usuarios
• Editar información de usuarios existentes

ROLES QUE PUEDE CREAR CADA TIPO DE ADMINISTRADOR:

Admin General:
• Puede crear: admin_dependencia, operador, consulta
• Puede seleccionar cualquier área del sistema
• Gestiona todos los usuarios del sistema

Admin Dependencia:
• Puede crear: operador, consulta
• Los usuarios heredan su misma área automáticamente
• Solo gestiona usuarios de su dependencia

CAMPOS PARA CREAR UN USUARIO:
• Nombre de usuario (único en el sistema)
• Contraseña (mínimo 6 caracteres)
• Nombre completo
• Correo electrónico
• Rol del usuario
• Permisos de eliminación (Sí/No)"""
    
    p = doc.add_paragraph()
    run = p.add_run(usuarios_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== CONSULTA PÚBLICA ==========
    print("Creando sección de consulta pública...")
    
    p = doc.add_paragraph()
    run = p.add_run("7. CONSULTA PÚBLICA DE ACTOS ADMINISTRATIVOS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    consulta_text = """El módulo de consulta pública permite a cualquier ciudadano buscar actos administrativos sin necesidad de iniciar sesión.

MÉTODOS DE BÚSQUEDA:
• Por número de cédula o NIT del contribuyente
• Por número de predio

RESULTADOS DE BÚSQUEDA:
• En Trámite: Actos administrativos vigentes
• Histórico: Actos administrativos archivados

INFORMACIÓN MOSTRADA:
• Número del acto administrativo
• Tipo de actuación
• Fecha de publicación
• Fecha de desfijación
• Organismo y área
• Opción de ver PDF asociado (si está disponible)

USO DE FILTROS EN HISTÓRICO:
• Filtrar por año de publicación
• Filtrar por dependencia/área
• Hacer clic en "Aplicar Filtros" para ver resultados"""
    
    p = doc.add_paragraph()
    run = p.add_run(consulta_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== SECCIÓN DE ROLES ==========
    print("Agregando sección de diferencias entre roles...")
    
    p = doc.add_paragraph()
    run = p.add_run("8. ROLES DE USUARIO Y SUS DIFERENCIAS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # ADMIN GENERAL
    p = doc.add_paragraph()
    run = p.add_run("ADMIN GENERAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(220, 38, 38)
    
    p = doc.add_paragraph()
    run = p.add_run("""El Administrador General tiene control total sobre el sistema. Este rol permite:
• Visualizar y gestionar todas las áreas y dependencias del sistema
• Crear usuarios de tipo admin_dependencia, operador y consulta en cualquier área
• Seleccionar el área al momento de crear nuevos usuarios
• Acceso completo a todos los registros sin importar el área de origen
• Visualización completa de la auditoría con filtros por cualquier usuario
• Gestión de todos los usuarios del sistema (activar, desactivar, editar)""")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ADMIN DEPENDENCIA
    p = doc.add_paragraph()
    run = p.add_run("ADMIN DEPENDENCIA")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(37, 99, 235)
    
    p = doc.add_paragraph()
    run = p.add_run("""El Administrador de Dependencia tiene control limitado a su área específica. Este rol:
• Solo puede ver y gestionar usuarios dentro de su dependencia
• Puede crear usuarios de tipo operador y consulta únicamente
• Los usuarios que crea heredan automáticamente su misma área
• Tiene acceso a registros de su organismo y áreas relacionadas
• Visualización de auditoría limitada a usuarios de su dependencia
• Gestión de usuarios solo de su área (activar, desactivar, editar)""")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # OPERADOR
    p = doc.add_paragraph()
    run = p.add_run("OPERADOR")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(22, 163, 74)
    
    p = doc.add_paragraph()
    run = p.add_run("""El Operador es el usuario encargado de la carga de datos. Este rol permite:
• Subir archivos Excel con datos de actos administrativos
• Visualizar únicamente sus propios registros cargados
• Eliminar registros solo si tiene permisos de eliminación habilitados
• Ver su propia actividad en la auditoría
• NO tiene acceso a la gestión de usuarios""")
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # CONSULTA
    p = doc.add_paragraph()
    run = p.add_run("CONSULTA")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(107, 114, 128)
    
    p = doc.add_paragraph()
    run = p.add_run("""El usuario de Consulta tiene acceso de solo lectura. Este rol permite:
• Consultar registros públicos de actos administrativos
• NO puede cargar archivos ni gestionar registros
• NO tiene acceso a la auditoría del sistema
• NO tiene acceso a la gestión de usuarios""")
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== SECCIÓN DE FTP ==========
    print("Agregando sección de subida de PDFs por FTP...")
    
    p = doc.add_paragraph()
    run = p.add_run("9. SUBIDA DE ARCHIVOS PDF VÍA FTP")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    ftp_text = """Los archivos PDF asociados a los actos administrativos deben subirse al servidor a través de FTP (File Transfer Protocol). A continuación se describe el proceso completo:

PASO 1: CONFIGURAR CLIENTE FTP
Utilice un cliente FTP como FileZilla, WinSCP o Cyberduck con las credenciales proporcionadas:
• Servidor/Host: Dirección del servidor FTP asignada por el administrador
• Usuario: Su usuario de FTP
• Contraseña: Su contraseña de FTP
• Puerto: Generalmente 21 (FTP estándar) o 22 (SFTP seguro)

PASO 2: NAVEGAR AL DIRECTORIO
Una vez conectado, navegue a la carpeta de destino: /uploads/pdfs/
O a su directorio personal de trabajo si tiene uno asignado.

PASO 3: SUBIR LOS ARCHIVOS PDF
• Arrastre los archivos PDF desde su computadora al directorio remoto
• El nombre del PDF debe corresponder con el nombre especificado en la columna del Excel
• Verifique que la transferencia se complete al 100%

PASO 4: VERIFICAR LA SUBIDA
• En el sistema web, vaya a la pestaña "Carga de Archivos"
• Haga clic en "Verificar PDFs" junto al archivo Excel pendiente
• El sistema mostrará si los PDFs fueron encontrados correctamente

NOMENCLATURA DE ARCHIVOS PDF:
El sistema busca los PDFs con los siguientes patrones:
• Nombre exacto especificado en la columna del archivo Excel
• Formato automático: archivo.xlsx → archivo.pdf
• El PDF puede tener extensión .xlsx.pdf o solo .pdf"""
    
    p = doc.add_paragraph()
    run = p.add_run(ftp_text)
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== PREGUNTAS FRECUENTES ==========
    print("Creando sección de preguntas frecuentes...")
    
    p = doc.add_paragraph()
    run = p.add_run("10. PREGUNTAS FRECUENTES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    faq_text = """¿Qué hago si olvidé mi contraseña?
Comuníquese con el administrador de su dependencia para que restablezca su contraseña.

¿Por qué no veo la pestaña de Gestión de Usuarios?
Esta pestaña solo está disponible para usuarios con rol admin_general o admin_dependencia.

¿Qué significa que un registro esté "eliminado"?
La eliminación es "lógica", lo que significa que el registro se oculta pero no se borra permanentemente. Un administrador puede restaurarlo si es necesario.

¿Por qué no puedo eliminar mis registros?
Debe tener habilitados los "permisos de eliminación" en su cuenta. Solicite a su administrador que le otorgue este permiso.

¿Cómo sé si mis PDFs se subieron correctamente?
En la pestaña "Carga de Archivos", haga clic en "Verificar PDFs" junto al archivo Excel. El sistema le indicará cuáles PDFs están disponibles y cuáles faltan.

¿Puedo modificar un registro después de cargarlo?
Actualmente el sistema no permite la edición directa de registros. Si necesita corregir datos, debe eliminar el registro y volver a cargarlo.

¿Cuánto tiempo permanecen los registros en "En Trámite"?
Los registros permanecen en estado "En Trámite" hasta que llegue su fecha de desfijación. Después de esta fecha, pasan automáticamente a "Histórico"."""
    
    p = doc.add_paragraph()
    run = p.add_run(faq_text)
    run.font.size = Pt(11)
    
    # Guardar documento
    print(f"\nGuardando documento: {output_path}")
    doc.save(output_path)
    
    print("\n" + "="*70)
    print("¡DOCUMENTO GENERADO EXITOSAMENTE!")
    print("="*70)
    print(f"Archivo creado: {output_path}")
    print("="*70)
    
    return output_path

if __name__ == "__main__":
    create_professional_guide()
