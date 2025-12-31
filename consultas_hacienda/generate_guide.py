from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_user_guide():
    document = Document()

    # Style configuration
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = document.add_heading('Manual de Usuario\nSistema de Consultas Hacienda', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n')
    
    intro = document.add_paragraph('Este documento sirve como guía detallada para el uso del Sistema de Consultas de Hacienda. El sistema permite la gestión de usuarios, carga de archivos y consulta de registros históricos y en trámite.')
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_page_break()

    # 1. Acceso al Sistema
    document.add_heading('1. Acceso al Sistema', level=1)
    document.add_paragraph('Para ingresar al sistema, diríjase a la pantalla de inicio de sesión. Ingrese su usuario y contraseña asignados.')
    
    # Image: Login
    img_login = r'C:/Users/15pab/.gemini/antigravity/brain/acb819a0-b797-4dc4-9cd0-cac1d4a7d78e/screenshot_login_1766951070082.png'
    if os.path.exists(img_login):
        document.add_picture(img_login, width=Inches(6))
    else:
        document.add_paragraph('[Error: Imagen de Login no encontrada]')

    document.add_paragraph('Roles disponibles:', style='List Bullet')
    document.add_paragraph('Admin General: Acceso total.', style='List Bullet')
    document.add_paragraph('Admin Dependencia: Gestión de su área.', style='List Bullet')
    document.add_paragraph('Operador: Carga y consulta.', style='List Bullet')
    document.add_paragraph('Consulta: Solo lectura.', style='List Bullet')

    # 2. Panel Principal (Dashboard)
    document.add_heading('2. Panel Principal (Consultas)', level=1)
    document.add_paragraph('Una vez dentro, visualizará el panel principal donde puede realizar búsquedas de expedientes y trámites.')
    
    # Image: Dashboard
    img_dash = r'C:/Users/15pab/.gemini/antigravity/brain/acb819a0-b797-4dc4-9cd0-cac1d4a7d78e/screenshot_dashboard_1766951140416.png'
    if os.path.exists(img_dash):
        document.add_picture(img_dash, width=Inches(6))

    # 3. Carga de Archivos
    document.add_heading('3. Carga de Archivos', level=1)
    document.add_paragraph('Los operadores y administradores pueden cargar nuevos registros al sistema desde la pestaña "Carga Masiva".')

    # Image: Upload
    img_upload = r'C:/Users/15pab/.gemini/antigravity/brain/acb819a0-b797-4dc4-9cd0-cac1d4a7d78e/screenshot_upload_1766951159196.png'
    if os.path.exists(img_upload):
        document.add_picture(img_upload, width=Inches(6))

    # 4. Gestión de Usuarios
    document.add_heading('4. Gestión de Usuarios', level=1)
    document.add_paragraph('Los administradores tienen acceso a la gestión de usuarios, donde pueden crear, editar y eliminar cuentas de su dependencia.')

    # Image: Users
    img_users = r'C:/Users/15pab/.gemini/antigravity/brain/acb819a0-b797-4dc4-9cd0-cac1d4a7d78e/screenshot_users_1766951180771.png'
    if os.path.exists(img_users):
        document.add_picture(img_users, width=Inches(6))

    # 5. Auditoría
    document.add_heading('5. Auditoría', level=1)
    document.add_paragraph('El sistema registra todas las acciones importantes. La pestaña "Auditoría" permite visualizar este historial.')

    # Image: Audit
    img_audit = r'C:/Users/15pab/.gemini/antigravity/brain/acb819a0-b797-4dc4-9cd0-cac1d4a7d78e/screenshot_audit_1766951203489.png'
    if os.path.exists(img_audit):
        document.add_picture(img_audit, width=Inches(6))

    # Save
    output_path = r'C:\Users\15pab\Desktop\consultas_hacienda\consultas_hacienda\GUIA_USUARIO_SISTEMA.docx'
    document.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_user_guide()
